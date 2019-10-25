# -*- coding:utf-8 -*-
#__author__:langzi
#__blog__:www.langzi.fun
import requests
import time
import aiohttp
import aiofiles
import aiomultiprocess
import multiprocessing
import asyncio
from binascii import b2a_hex
from string import whitespace
import urllib
import requests
requests.packages.urllib3.disable_warnings()
import re
from urllib.parse import urlparse,urlencode,parse_qs,parse_qsl,quote_plus
from urllib.request import urlopen
import mechanize
import http.client
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor
import random
import sys
import os
import time
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches
import multiprocessing
import re
import random
import requests
from concurrent.futures import ProcessPoolExecutor
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab
# 输入屏幕左上角和右下角的坐标
from urllib.parse import urlparse
import contextlib
import configparser
import pymysql
import datetime
import os
import time
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches
import multiprocessing
import re
import random
import requests
from concurrent.futures import ProcessPoolExecutor
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from PIL import ImageGrab
# 输入屏幕左上角和右下角的坐标
from urllib.parse import urlparse

import threading

lock1 = threading.Lock()
lock2 = threading.Lock()


cfg = configparser.ConfigParser()
cfg.read('../Config.ini')
user = cfg.get("Server", "username")
passwd = cfg.get("Server", "password")
host = cfg.get("Server", "host")
Dbname = cfg.get("Server", "db")
port = int(cfg.get("Server", "port"))

@contextlib.contextmanager
def connect_mysql():
    coon = pymysql.connect(user=user, passwd=passwd, host=host, db=Dbname, port=port, charset='utf8')
    cursor = coon.cursor()
    try:
        yield cursor
    except Exception as e:
        print(e)
        pass
    finally:
        coon.commit()
        cursor.close()
        coon.close()


headerss = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

def writedata(x):
    with open('log.txt', 'a+')as aa:
        aa.write('***********************************' + '\n')
        aa.write(str(time.strftime('%Y-%m-%d:%H:%M:%S   ', time.localtime())) + str(x) + '\n')


class Get_Info:
    def __init__(self,url):
        self.url=url
        self.title_parrten = 'class="w61-0"><div class="ball">(.*?)</div></td>'  # group(1) 正常
        self.ip_parrten = '>IP：(.*?)</a></div>'  # group(1) 正常
        self.ages = '" target="_blank">(.*?)</a></div></div>'  # group(1)
        self.whois_id = '备案号：</span><a href=.*?" target="_blank">(.*?)</a></div>'  # 需group(1)
        self.whois_type = '<span>性质：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_name = '<span>名称：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_time = '<span>审核时间：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.include_baidu = '<div class="Ma01LiRow w12-1 ">(.*?)</div>'  # group(1)
        self.infos = '<div class="MaLi03Row w180">(.*?)</div>'  # 要findall 0，1，2，3
        self.result={'百度权重':'',
                    '网站主页':'',
                     '网站描述':'',
                   '网站标题':'',
                   'IP__坐标':'',
                   '所属__IP':'',
                   '网站年龄':'',
                   '备案编号':'',
                   '备案性质':'',
                   '备案名称':'',
                   '备案时间':'',
                   '百度收录':'',
                   '协议类型':'',
                   '页面类型':'',
                   '服务类型':'',
                   '程序语言':''}

    def get_baidu_weight(self):
        time.sleep(random.randint(1, 5))
        x = str(random.randint(1, 9))
        data = {
            't': 'rankall',
            'on': 1,
            'type': 'baidupc',
            'callback': 'jQuery111303146901980779846_154444474116%s' % (x),
            'host': self.url
        }

        headers = {

            'Accept': 'text/javascript, application/javascript, application/ecmascript, application/x-ecmascript, */*; q=0.01',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'Cookie': 'UM_distinctid=165af67ee6f352-07238a34ed3941-9393265-1fa400-165af67ee70473; CNZZDATA5082706=cnzz_eid%3D832961605-1544438317-null%26ntime%3D1544443717; Hm_lvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; Hm_lpvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; qHistory=aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9iYWlkdW1vYmlsZS8r55m+5bqm56e75Yqo5p2D6YeNfGh0dHA6Ly9yYW5rLmNoaW5hei5jb20vcmFua2FsbC8r5p2D6YeN57u85ZCI5p+l6K+ifGh0dHA6Ly9yYW5rLmNoaW5hei5jb20r55m+5bqm5p2D6YeN5p+l6K+ifGh0dHA6Ly9pbmRleC5jaGluYXouY29tLyvlhbPplK7or43lhajnvZHmjIfmlbB8aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9yYW5rL2hpc3RvcnkuYXNweCvmnYPph43ljoblj7Lmn6Xor6I=',
            'Host': 'rank.chinaz.com',
            'Origin': 'http://rank.chinaz.com',
            'Referer': 'http://rank.chinaz.com',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest'

        }
        try:
            urls = 'http://rank.chinaz.com/ajaxseo.aspx?t=rankall&on=1&type=undefined&callback=jQuery111303146901980779846_154444474116%s' % (
                x)

            r = requests.post(url=urls, headers=headers, data=data)
            try:
                res = re.search(b',"br":(\d),"beforBr', r.content).group(1)
            except:
                pass
            if res:
                return res.decode()
            else:
                return '0'
        except:
            return '获取失败'


    def get_info_from_pattren(self,pattren, result):
        try:
            res = re.search(pattren, result).group(1)
            return res
        # return str(res.encode('utf-8'))
        except:
            return '暂无信息'

    def scan_seo(self):
        UA = random.choice(headerss)
        headers = {'User-Agent': UA}
        domain_url = self.url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        urls = 'http://seo.chinaz.com/' + domain_url
        # url,title,weights,ip,ages,whois_id,whois_type,whois_name,whois_time
        # 网址，标题，百度权重，ip信息，年龄，备案号，备案性质，备案名称，备案时间
        # include_baidu,request,text,service,language
        # 百度收录，，协议类型，页面类型，服务器类型，程序语言
        try:
            req = requests.get(urls, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            r = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))

        try:
            req = requests.get(self.url, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            rss = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))


        self.result['网站描述'] = '暂无信息'
        try:
            description = re.search('description.*?content=(.*?)>',rss,re.S).group(1)
            self.result['网站描述'] = description.strip()
        except:
            pass

        self.result['百度权重'] = str(self.get_baidu_weight())

        self.result['网站主页'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]

        try:
            req1 = requests.get(url=self.url, headers=headers, verify=False, timeout=10)
            encoding = requests.utils.get_encodings_from_content(req1.text)[0]
            rress = req1.content.decode(encoding, 'replace')
            title_pattern = '<title>(.*?)</title>'
            title = re.search(title_pattern, rress, re.S | re.I)
            self.result['网站标题'] = str(title.group(1))
        except:
            self.result['网站标题'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]


        ip_infos = self.get_info_from_pattren(self.ip_parrten, r)
        if '[' in ip_infos:
            ip, address = ip_infos.split('[')[0], ip_infos.split('[')[1]
            self.result['IP__坐标'] = address.replace(']', '')
            self.result['所属__IP'] = ip
        else:
            self.result['IP__坐标'] = '获取失败'
            self.result['所属__IP'] = '获取失败'

        self.result['网站年龄'] = self.get_info_from_pattren(self.ages, r)
        self.result['备案编号'] = self.get_info_from_pattren(self.whois_id, r)
        self.result['备案性质'] = self.get_info_from_pattren(self.whois_type, r)
        self.result['备案名称'] = self.get_info_from_pattren(self.whois_name, r)
        self.result['备案时间'] = self.get_info_from_pattren(self.whois_time, r)
        self.result['百度收录'] = self.get_info_from_pattren(self.include_baidu, r)

        dd = re.findall(self.infos, r, re.S)
        resu = ['暂无信息' if x.replace(' ', '') is '' else x for x in dd]
        try:
            self.result['协议类型'] = resu[0]
        except:
            self.result['协议类型'] = '获取失败'

        try:
            self.result['页面类型'] = resu[1]
        except:
            self.result['页面类型'] = '获取失败'

        try:
            self.result['服务类型'] = resu[2]
        except:
            self.result['服务类型'] = '获取失败'

        try:
            self.result['程序语言'] = resu[3]
        except:
            self.result['程序语言'] = '获取失败'
        return self.result





def mains(url,org):
    print('Check : {}'.format(url))

    print('开始保存文档....')
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    DaHei = document.styles.add_style('DaHei', 1)
    # 设置字体尺寸
    DaHei.font.size = Pt(16)
    # 设置字体颜色
    DaHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    DaHei.font.name = '仿宋'
    DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHei = document.styles.add_style('ZhongHei', 1)
    # 设置字体尺寸
    ZhongHei.font.size = Pt(10)
    # 设置字体颜色
    ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
    # 黑色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHei.font.name = '仿宋'
    ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    ZhongHOng = document.styles.add_style('ZhongHOng', 1)
    # 设置字体尺寸
    ZhongHOng.font.size = Pt(9)
    # 设置字体颜色
    ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
    # 红色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    ZhongHOng.font.name = '仿宋'
    ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    XiaoLv = document.styles.add_style('XiaoLv', 1)
    # 设置字体尺寸
    XiaoLv.font.size = Pt(6.5)
    # 设置字体颜色
    XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
    # 绿色
    # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    XiaoLv.font.name = '仿宋'
    XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    domain = urlparse(url).netloc
    try:
        filename = url + '__存在CORS劫持漏洞.docx'
    except:
        pass
    finally:
        filename = filename.replace('__','_'+str(random.randint(1,2000))+'_')
        filename = '../result/' + filename

    document.add_paragraph('基本信息', 'Title')
    document.add_paragraph('漏洞类型：CORS劫持漏洞', 'Subtitle')
    document.add_paragraph('漏洞等级：低危', 'Subtitle')
    document.add_paragraph('厂商信息：此处需要手动搜索', 'Subtitle')
    document.add_page_break()

    document.add_paragraph('漏洞简述', 'Title')
    document.add_paragraph('漏洞描述', 'Subtitle')
    document.add_paragraph('可以造成对该网址的数据通过跨源资源共享方式进行劫持获取', 'DaHei')
    document.add_paragraph('产生原因', 'Subtitle')
    document.add_paragraph('''
    1 CORS是W3C出的一个标准，其思想是使用自定义的HTTP头部让浏览器与服务器进行沟通。
    2 因为开发者需要进行跨域进行获取资源，所以如果没有对请求资源域进行校验检测，则会出现黑客伪装请求，进行劫持数据
    
    ''', 'DaHei')
    document.add_paragraph('漏洞危害', 'Subtitle')
    document.add_paragraph(
       '''
    1 造成跨域问题，可能造成信息泄露等危害。
    2 可能造成数据劫持，泄露敏感数据
    
       ''','DaHei')

    document.add_page_break()

    document.add_paragraph('漏洞详情', 'Title')
    document.add_paragraph('网站漏洞报表', 'Subtitle')
    document.add_paragraph('漏洞地址:'+url,'ZhongHei')
    document.add_paragraph('Origin对象:'+org,'ZhongHei')
    urls = 'http://'+url
    try:
        document.add_paragraph('网站信息报表', 'Subtitle')
        info = Get_Info(urls)
        infos = info.scan_seo()
        for k,v in infos.items():
            document.add_paragraph(k+'  '+v,'ZhongHei')
        document.add_page_break()
    except Exception as e:
        writedata(str(e))


    #document.add_paragraph('网站漏洞复现(使用Firefox浏览器)', 'Subtitle')
    #document.add_paragraph('下载链接','DaHei')
    #document.add_paragraph(url,'ZhongHOng')
    #document.add_paragraph('返回结果','DaHei')
    #document.add_picture('image/code.png',width=Inches(7))
    #document.add_paragraph('复现完成，验证存在源码泄露漏洞','DaHei')

    document.add_page_break()
    document.add_paragraph('修复建议', 'Title')
    document.add_paragraph('根源上进行修复', 'Subtitle')
    document.add_paragraph('''
    在防御过程中，存在很多配置不当，或者配置错误的情况，例如使用*,或者使用多个域名list,或者使用null,或者nginx配置规则存在错误等等。
    ''','DaHei')
    try:
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Report Powered By Langzi"
    except:
        pass
    document.save(filename)

headers={'User-Agent': 'Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3', 'Accept': 'Accept:text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8', 'Cache-Control': 'max-age=0', 'referer': 'http://www.soso.com', 'Accept-Charset': 'GBK,utf-8;q=0.7,*;q=0.3'}

def main(url):
    domain = url.replace('https://','').replace('http://','').replace('www.','')
    if 'http://' or 'https://' not in domain:
        domain = 'http://' + domain.strip()
    print('scan:'+domain)

    try:
        characters = '!@#$%^&*()_+~/*'
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.157 Safari/537.36',
            'Origin': '{}'.format('http://www.baidu.com')
        }
        req1 = requests.get(domain, headers=headers, timeout=(5, 20), verify=False, allow_redirects=False)
        if "Access-Control-Allow-Origin" and "Access-Control-Allow-Credentials" in req1.headers:
            if req1.headers['Access-Control-Allow-Origin'] == headers['Origin']:
                with open('cors_result.txt', 'a+') as f:
                    f.write(
                        '{} {} {} \n'.format(domain.replace('http://', ''), req1.headers['Access-Control-Allow-Origin'],
                                             req1.headers['Access-Control-Allow-Credentials']))
            else:
                headers['Origin'] = domain + '.baidu.com'
                req2 = requests.get(domain, headers=headers, timeout=(5, 20), verify=False, allow_redirects=False)
                if req2.headers['Access-Control-Allow-Origin'] == headers['Origin']:
                    with open('cors_result.txt', 'a+') as f:
                        f.write('{} {} {} \n '.format(domain.replace('http://', ''),
                                                      req2.headers['Access-Control-Allow-Origin'],
                                                      req2.headers['Access-Control-Allow-Credentials']))
                else:
                    for character in characters:
                        headers['Origin'] = domain + character + '.baidu.com'
                        req3 = requests.get(domain, headers=headers, timeout=(5, 20), verify=False,
                                            allow_redirects=False)
                        if req3.headers['Access-Control-Allow-Origin'] == headers['Origin']:
                            
                            with open('cors_result.txt', 'a+') as f:
                                f.write('{} {} {} \n').format(domain.replace('http://', ''),
                                                              req3.headers['Access-Control-Allow-Origin'],
                                                              req3.headers['Access-Control-Allow-Credentials'])
                        else:
                            pass

    except Exception as e:
        pass


if __name__ == '__main__':
    multiprocessing.freeze_support()
    if os.path.exists('cors_result.txt'):
        os.remove('cors_result.txt')

    if os.path.exists('result'):
        pass
    else:
        os.mkdir('result')

    if os.path.exists('image'):
        pass
    else:
        os.mkdir('image')
    print('''
         _____  __    __  _____   _____        ___   _____   _____  
        | ____| \ \  / / |_   _| |  _  \      /   | /  ___| |_   _| 
        | |__    \ \/ /    | |   | |_| |     / /| | | |       | |   
        |  __|    }  {     | |   |  _  /    / / | | | |       | |   
        | |___   / /\ \    | |   | | \ \   / /  | | | |___    | |   
        |_____| /_/  \_\   |_|   |_|  \_\ /_/   |_| \_____|   |_|   

    ''')

    while 1:
        with connect_mysql() as conn:
            sql1 = 'select count(id) from Sec_Index where cors=0'
            conn.execute(sql1)
            res = int(conn.fetchone()[0])
            if res == 0:
                print('数据库全部URL对cors扫描完毕')
                os.system('pause')
            if res > 100:
                sql2 = 'select url from Sec_Index where cors=0 limit 100'
                sql3 = 'update Sec_Index set cors=1 where cors=0 limit 100'
                conn.execute(sql2)
                result = [x[0] for x in (conn.fetchall())]
                print('任务计划:'+str(result))
                conn.execute(sql3)
                for u in result:
                    main(u)
            else:
                sql2 = 'select url from Sec_Index where cors=0 limit '+str(res)
                sql3 = 'update Sec_Index set cors=1 where cors=0 limit '+str(res)
                result = [x[0] for x in (conn.fetchall())]
                print('任务计划:'+str(result))
                conn.execute(sql3)
                for u in result:
                    main(u)
        if os.path.exists('cors_result.txt'):
            New_start = 'cors_result.txt'
            list_ = [x.strip() for x in open(New_start, 'r', encoding='utf-8').readlines()]
            for x in list_:
                print('发现漏洞:'+x)
                url,org = x.split(' ')[0],x.split(' ')[1]
                mains(url=url,org=org)
            os.remove('cors_result.txt')

