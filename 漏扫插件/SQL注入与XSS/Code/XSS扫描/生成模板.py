# -*- coding:utf-8 -*-
#__author__:langzi
#__blog__:www.langzi.fun

import sys
reload(sys)
sys.setdefaultencoding('utf-8')
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches

from urlparse import urlparse

def main(url,keyword,method,payload):
        document = Document()
        DaHei = document.styles.add_style('DaHei', 1)
        # 设置字体尺寸
        DaHei.font.size = Pt(16)
        # 设置字体颜色
        DaHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        DaHei.font.name = '仿宋'.decode('utf-8')
        DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

        ZhongHei = document.styles.add_style('ZhongHei', 1)
        # 设置字体尺寸
        ZhongHei.font.size = Pt(10)
        # 设置字体颜色
        ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
        # 黑色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHei.font.name = '仿宋'.decode('utf-8')
        ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

        ZhongHOng = document.styles.add_style('ZhongHOng', 1)
        # 设置字体尺寸
        ZhongHOng.font.size = Pt(9)
        # 设置字体颜色
        ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
        # 红色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        ZhongHOng.font.name = '仿宋'.decode('utf-8')
        ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

        XiaoLv = document.styles.add_style('XiaoLv', 1)
        # 设置字体尺寸
        XiaoLv.font.size = Pt(6.5)
        # 设置字体颜色
        XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
        # 绿色
        # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
        # 居中文本
        XiaoLv.font.name = '仿宋'.decode('utf-8')
        XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))
        domain = urlparse(url).netloc
        filename = url.split('//')[1].split('/')[0] + '__页面存在反射型XSS漏洞.docx'
        document.add_paragraph('基本信息'.decode('utf-8'), 'Title')
        document.add_paragraph('漏洞类型：XSS跨站注入漏洞'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('漏洞等级：中危'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('厂商信息：此处需要手动搜索'.decode('utf-8'), 'Subtitle')
        document.add_page_break()

        document.add_paragraph('漏洞简述'.decode('utf-8'), 'Title')
        document.add_paragraph('漏洞描述'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('跨站脚本攻击的英文全称是Cross Site Script，为了和样式表区分，缩写为XSS。发生的原因是网站将用户输入的内容输出到页面上，在这个过程中可能有恶意代码被浏览器执行。跨站脚本攻击,它指的是恶意攻击者往Web页面里插入恶意html代码，当用户浏览该页之时，嵌入其中Web里面的html代码会被执行，从而达到恶意用户的特殊目的。'.decode('utf-8'), 'DaHei')
        document.add_paragraph('漏洞危害'.decode('utf-8'), 'Subtitle')
        document.add_paragraph(
           '''
        1.窃取cookies，读取目标网站的cookie发送到黑客的服务器上
        2.读取用户未公开的资料，如果：邮件列表或者内容、系统的客户资料，联系人列表等等。

           '''.decode('utf-8'),'DaHei')

        document.add_page_break()

        document.add_paragraph('漏洞详情'.decode('utf-8'), 'Title')
        document.add_paragraph('网站漏洞报表'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('注入网址:'.decode('utf-8')+url,'ZhongHei')
        document.add_paragraph('漏洞参数:'.decode('utf-8')+keyword,'ZhongHei')
        document.add_paragraph('请求方式:'.decode('utf-8')+method,'ZhongHei')
        document.add_paragraph('攻击载荷:'.decode('utf-8')+payload,'ZhongHei')
        document.add_page_break()


        document.add_paragraph('网站漏洞复现(使用Firefox浏览器)'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('访问链接'.decode('utf-8'),'DaHei')
        document.add_paragraph(url,'ZhongHOng')
        document.add_paragraph('返回结果'.decode('utf-8'),'DaHei')
        document.add_picture('code.png',width=Inches(7))
        document.add_paragraph('复现完成，验证存在反射型XSS漏洞'.decode('utf-8'),'DaHei')

        document.add_page_break()
        document.add_paragraph('修复建议'.decode('utf-8'), 'Title')
        document.add_paragraph('代码原理防护'.decode('utf-8'), 'Subtitle')
        document.add_paragraph('''
        1、	输入验证：某个数据被接受为可被显示或存储之前，使用标准输入验证机制，验证所有输入数据的长度、类型、语法以及业务规则。
        2、	输出编码：数据输出前，确保用户提交的数据已被正确进行entity编码，建议对所有字符进行编码而不仅局限于某个子集。
        3、	明确指定输出的编码方式：不要允许攻击者为你的用户选择编码方式(如ISO 8859-1或 UTF 8)。
        4、	注意黑名单验证方式的局限性：仅仅查找或替换一些字符(如"<" ">"或类似"script"的关键字)，很容易被XSS变种攻击绕过验证机制。
        5、	警惕规范化错误：验证输入之前，必须进行解码及规范化以符合应用程序当前的内部表示方法。请确定应用程序对同一输入不做两次解码。对客户端提交的数据进行过滤，一般建议过滤掉双引号（”）、尖括号（<、>）等特殊字符，或者对客户端提交的数据中包含的特殊字符进行实体转换，比如将双引号（”）转换成其实体形式&quot;，<对应的实体形式是&lt;，<对应的实体形式是&gt;
        '''.decode('utf-8'),'DaHei')
        document.save(filename.decode('utf-8'))
main('ttp://127.0.0.1/xss/level1.php?name=1','name','GET','<script>prompt(1)</script>')
